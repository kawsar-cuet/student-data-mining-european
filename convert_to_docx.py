#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Simple LaTeX to DOCX Converter
Converts JOURNAL_METHODOLOGY.tex to Word format
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_latex_text(text):
    """Remove LaTeX commands and clean text"""
    # Remove comments
    text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)
    
    # Remove common LaTeX commands
    text = re.sub(r'\\textbf\{([^}]+)\}', r'**\1**', text)
    text = re.sub(r'\\textit\{([^}]+)\}', r'*\1*', text)
    text = re.sub(r'\\emph\{([^}]+)\}', r'*\1*', text)
    text = re.sub(r'\\cite\{([^}]+)\}', r'[\1]', text)
    text = re.sub(r'\\citep\{([^}]+)\}', r'[\1]', text)
    text = re.sub(r'\\citet\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\ref\{([^}]+)\}', r'[REF:\1]', text)
    text = re.sub(r'\\label\{([^}]+)\}', '', text)
    
    # Remove other common commands
    text = re.sub(r'\\(?:texttt|textsc|textcolor\{[^}]+\})\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+(?:\[[^\]]*\])?\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+\s+', ' ', text)
    
    # Clean up special characters
    text = text.replace('~', ' ')
    text = text.replace('---', '—')
    text = text.replace('--', '–')
    
    return text.strip()

def convert_latex_to_docx(input_file, output_file):
    """Convert LaTeX file to DOCX"""
    print(f"Converting {input_file} to {output_file}...")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Extract title
    title_match = re.search(r'\\title\{([^}]+)\}', content)
    if title_match:
        title = clean_latex_text(title_match.group(1))
        p = doc.add_heading(title, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Extract abstract
    abstract_match = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', content, re.DOTALL)
    if abstract_match:
        doc.add_heading('Abstract', 1)
        abstract = clean_latex_text(abstract_match.group(1))
        doc.add_paragraph(abstract)
    
    # Extract sections
    sections = re.findall(r'\\section\{([^}]+)\}(.*?)(?=\\section\{|\\end\{document\}|$)', content, re.DOTALL)
    
    for section_title, section_content in sections:
        section_title = clean_latex_text(section_title)
        doc.add_heading(section_title, 1)
        
        # Extract subsections
        subsections = re.findall(r'\\subsection\{([^}]+)\}(.*?)(?=\\subsection\{|\\section\{|$)', section_content, re.DOTALL)
        
        if subsections:
            for subsection_title, subsection_content in subsections:
                subsection_title = clean_latex_text(subsection_title)
                doc.add_heading(subsection_title, 2)
                
                # Clean content
                text = clean_latex_text(subsection_content)
                # Split by paragraphs
                paragraphs = text.split('\n\n')
                for para in paragraphs:
                    para = para.strip()
                    if para and not para.startswith('\\'):
                        doc.add_paragraph(para)
        else:
            # No subsections, just add content
            text = clean_latex_text(section_content)
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if para and not para.startswith('\\'):
                    doc.add_paragraph(para)
    
    # Save document
    doc.save(output_file)
    print(f"✓ Conversion complete: {output_file}")
    print(f"  File size: {os.path.getsize(output_file) / 1024:.1f} KB")

if __name__ == '__main__':
    input_tex = 'docs/JOURNAL_METHODOLOGY.tex'
    output_docx = 'docs/JOURNAL_METHODOLOGY.docx'
    
    try:
        convert_latex_to_docx(input_tex, output_docx)
    except Exception as e:
        print(f"Error: {e}")
        print("\nNote: This script requires python-docx package.")
        print("Install it with: pip install python-docx")
